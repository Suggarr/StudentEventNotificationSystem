
from PyQt6 import QtCore, QtGui, QtWidgets
from PyQt6.QtWidgets import QTableWidgetItem
import Add_Student
from PyQt6.QtWidgets import QMessageBox
from docx import Document
from docxtpl import DocxTemplate
import sys
import easygui



class Ui_MainForm(object):
    def __init__(self):
        self.marks=[0,0,0,0,0,0,0,0,0,0]
        self.count = [0, 0, 0]  # количество студентов, присутсувющих и нет
        self.gpa = [0, 0]
        self.sortspisok = []
        self.inf=[]
    def setupUi(self, MainForm, information):
        self.inf=information
        MainForm.setObjectName("MainForm")
        MainForm.resize(700, 600)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Policy.Fixed, QtWidgets.QSizePolicy.Policy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(MainForm.sizePolicy().hasHeightForWidth())
        MainForm.setSizePolicy(sizePolicy)
        MainForm.setMinimumSize(QtCore.QSize(700, 600))
        MainForm.setMaximumSize(QtCore.QSize(700, 600))
        MainForm.setContextMenuPolicy(QtCore.Qt.ContextMenuPolicy.NoContextMenu)
        icon = QtGui.QIcon()
        icon.addPixmap(QtGui.QPixmap("resources/logo.png"), QtGui.QIcon.Mode.Normal, QtGui.QIcon.State.On)
        MainForm.setWindowIcon(icon)
        MainForm.setLayoutDirection(QtCore.Qt.LayoutDirection.LeftToRight)
        MainForm.setAutoFillBackground(False)
        _translate = QtCore.QCoreApplication.translate
        MainForm.setWindowTitle(_translate("MainForm", "ExamList"))
        self.centralwidget = QtWidgets.QWidget(parent=MainForm)
        self.centralwidget.setObjectName("centralwidget")
        self.tableWidget = QtWidgets.QTableWidget(parent=self.centralwidget)
        self.tableWidget.setGeometry(QtCore.QRect(0, 170, 701, 280))
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Policy.Minimum, QtWidgets.QSizePolicy.Policy.Minimum)
        sizePolicy.setHorizontalStretch(100)
        sizePolicy.setVerticalStretch(100)
        sizePolicy.setHeightForWidth(self.tableWidget.sizePolicy().hasHeightForWidth())
        self.tableWidget.setSizePolicy(sizePolicy)
        self.tableWidget.setRowCount(2)
        self.tableWidget.setColumnCount(5)
        self.tableWidget.setObjectName("tableWidget")
        item = QtWidgets.QTableWidgetItem()
        self.tableWidget.setHorizontalHeaderItem(0, item)
        item = QtWidgets.QTableWidgetItem()
        item.setTextAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        item.setFlags(QtCore.Qt.ItemFlag.ItemIsEnabled)
        self.tableWidget.setItem(0, 0, item)
        item = QtWidgets.QTableWidgetItem()
        item.setTextAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        item.setFlags(QtCore.Qt.ItemFlag.ItemIsEnabled)
        self.tableWidget.setItem(0, 1, item)
        item = QtWidgets.QTableWidgetItem()
        item.setTextAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        item.setFlags(QtCore.Qt.ItemFlag.ItemIsEnabled)
        self.tableWidget.setItem(0, 2, item)
        item = QtWidgets.QTableWidgetItem()
        item.setTextAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        item.setFlags(QtCore.Qt.ItemFlag.ItemIsEnabled)
        self.tableWidget.setItem(0, 3, item)
        item = QtWidgets.QTableWidgetItem()
        item.setTextAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        item.setFlags(QtCore.Qt.ItemFlag.ItemIsEnabled)
        self.tableWidget.setItem(0, 4, item)
        self.tableWidget.horizontalHeader().setVisible(False)
        self.tableWidget.horizontalHeader().setDefaultSectionSize(166)
        self.tableWidget.verticalHeader().setVisible(False)
        self.tableWidget.verticalHeader().setHighlightSections(True)
        self.tableWidget.verticalHeader().setSortIndicatorShown(False)
        self.textEdit = QtWidgets.QTextEdit(parent=self.centralwidget)
        self.textEdit.setGeometry(QtCore.QRect(0, 0, 700, 170))
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Policy.Minimum, QtWidgets.QSizePolicy.Policy.Minimum)
        sizePolicy.setHorizontalStretch(100)
        sizePolicy.setVerticalStretch(100)
        sizePolicy.setHeightForWidth(self.textEdit.sizePolicy().hasHeightForWidth())
        self.textEdit.setSizePolicy(sizePolicy)
        self.textEdit.setMinimumSize(QtCore.QSize(100, 100))
        self.textEdit.setBaseSize(QtCore.QSize(50, 50))
        self.textEdit.setFocusPolicy(QtCore.Qt.FocusPolicy.StrongFocus)
        self.textEdit.setContextMenuPolicy(QtCore.Qt.ContextMenuPolicy.NoContextMenu)
        self.textEdit.setLayoutDirection(QtCore.Qt.LayoutDirection.LeftToRight)
        self.textEdit.setVerticalScrollBarPolicy(QtCore.Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.textEdit.setHorizontalScrollBarPolicy(QtCore.Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.textEdit.setTextInteractionFlags(QtCore.Qt.TextInteractionFlag.TextSelectableByKeyboard|QtCore.Qt.TextInteractionFlag.TextSelectableByMouse)
        self.textEdit.setObjectName("textEdit")
        self.textEdit.setReadOnly(True)
        self.textEdit_2 = QtWidgets.QTextEdit(parent=self.centralwidget)
        self.textEdit_2.setGeometry(QtCore.QRect(0, 450, 821, 151))
        self.textEdit_2.setMinimumSize(QtCore.QSize(821, 151))
        self.textEdit_2.setMaximumSize(QtCore.QSize(821, 151))
        self.textEdit_2.setFocusPolicy(QtCore.Qt.FocusPolicy.NoFocus)
        self.textEdit_2.setObjectName("textEdit_2")
        self.textEdit_2.setReadOnly(True)
        MainForm.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(parent=MainForm)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 819, 26))
        self.menubar.setObjectName("menubar")
        self.menu = QtWidgets.QMenu(parent=self.menubar)
        self.menu.setObjectName("menu")
        self.menu_3 = QtWidgets.QMenu(parent=self.menubar)
        self.menu_3.setObjectName("menu_3")
        self.menu_2 = QtWidgets.QMenu(parent=self.menubar)
        self.menu_2.setObjectName("menu_2")
        MainForm.setMenuBar(self.menubar)
        self.create = QtGui.QAction(parent=MainForm)
        self.create.setObjectName("create")
        self.load = QtGui.QAction(parent=MainForm)
        self.load.setObjectName("load")
        self.save = QtGui.QAction(parent=MainForm)
        self.save.setObjectName("save")
        self.export = QtGui.QAction(parent=MainForm)
        self.export.setObjectName("export")
        self.sending = QtGui.QAction(parent=MainForm)
        self.sending.setObjectName("sending")
        self.add_student = QtGui.QAction(parent=MainForm)
        self.add_student.setObjectName("add_student")
        self.change_student = QtGui.QAction(parent=MainForm)
        self.change_student.setObjectName("change_student")
        self.delete_student = QtGui.QAction(parent=MainForm)
        self.delete_student.setObjectName("delete_stuent")
        self.user_inf = QtGui.QAction(parent=MainForm)
        self.user_inf.setObjectName("action")
        self.prog_inf = QtGui.QAction(parent=MainForm)
        self.prog_inf.setObjectName("action_3")
        self.menu.addAction(self.create)
        self.menu.addAction(self.load)
        self.menu.addAction(self.save)
        self.menu.addAction(self.export)
        self.menu.addAction(self.sending)
        self.menu_3.addAction(self.add_student)
        self.menu_3.addAction(self.change_student)
        self.menu_3.addAction(self.delete_student)
        self.menu_2.addAction(self.user_inf)
        self.menu_2.addAction(self.prog_inf)
        self.menubar.addAction(self.menu.menuAction())
        self.menubar.addAction(self.menu_3.menuAction())
        self.menubar.addAction(self.menu_2.menuAction())
        self.tableWidget.resizeRowsToContents()
        self.tableWidget.resizeColumnToContents(0)
        self.tableWidget.setSpan(0, 0, 2, 1)
        self.tableWidget.setSpan(0, 1, 2, 1)
        self.tableWidget.setSpan(0, 2, 2, 1)
        self.tableWidget.setSpan(0, 3, 2, 1)
        self.tableWidget.setSpan(0, 4, 2, 1)
        self.Add = QtWidgets.QPushButton(parent=MainForm)#начало создания скрытых кнопок и т.д.
        self.Add.setGeometry(QtCore.QRect(305, 560, 90, 40))
        self.Add.setStyleSheet("background-color: qlineargradient(spread:pad, x1:0.492537, y1:0.807, x2:1, y2:0, stop:0 rgba(186, 186, 186, 255), stop:1 rgba(255, 255, 255, 255));")
        self.Add.setAutoDefault(False)
        self.Add.setDefault(False)
        self.Add.setFlat(False)
        self.Add.setObjectName("Add")
        self.Add.setVisible(False)
        self.Cancel = QtWidgets.QPushButton(parent=MainForm)
        self.Cancel.setGeometry(QtCore.QRect(610, 560, 90, 40))
        self.Cancel.setStyleSheet("background-color: qlineargradient(spread:pad, x1:0.492537, y1:0.807, x2:1, y2:0, stop:0 rgba(186, 186, 186, 255), stop:1 rgba(255, 255, 255, 255));")
        self.Cancel.setAutoDefault(False)
        self.Cancel.setDefault(False)
        self.Cancel.setFlat(False)
        self.Cancel.setObjectName("Cancel")
        self.Cancel.setVisible(False)

        self.label = QtWidgets.QLabel(parent=self.centralwidget)
        self.label.setObjectName("label")
        self.label.setVisible(False)
        self.label_2 = QtWidgets.QLabel(parent=self.centralwidget)
        self.label_2.setObjectName("label_2")
        self.label_2.setVisible(False)
        self.label_3 = QtWidgets.QLabel(parent=self.centralwidget)
        self.label_3.setObjectName("label_3")
        self.label_3.setVisible(False)
        self.secondnames = QtWidgets.QTextEdit(parent=self.centralwidget)
        self.secondnames.setObjectName("secondnames")
        self.secondnames.setVisible(False)
        self.number_book = QtWidgets.QTextEdit(parent=self.centralwidget)
        self.number_book.setObjectName("number_book")
        self.number_book.setVisible(False)
        self.ball_int = QtWidgets.QComboBox(parent=self.centralwidget)
        self.ball_int.setObjectName("ball_int")
        self.ball_int.addItem("")
        self.ball_int.addItem("")
        self.ball_int.addItem("")
        self.ball_int.addItem("")
        self.ball_int.addItem("")
        self.ball_int.addItem("")
        self.ball_int.addItem("")
        self.ball_int.addItem("")
        self.ball_int.addItem("")
        self.ball_int.addItem("")
        self.ball_int.addItem("")
        self.ball_int.addItem("")
        self.ball_int.setVisible(False)
        self.label_del = QtWidgets.QLabel(parent=self.centralwidget)#Удаление
        self.label_del.setObjectName("label_del")
        self.label_del.setVisible(False)
        self.number_stud = QtWidgets.QComboBox(parent=self.centralwidget)
        self.number_stud.setObjectName("number_stud")
        self.number_stud.setVisible(False)
        self.delete_button = QtWidgets.QPushButton(parent=MainForm)  # начало создания скрытых кнопок и т.д.
        self.delete_button.setGeometry(QtCore.QRect(287, 560, 90, 40))
        self.delete_button.setStyleSheet("background-color: qlineargradient(spread:pad, x1:0.492537, y1:0.807, x2:1, y2:0, stop:0 rgba(186, 186, 186, 255), stop:1 rgba(255, 255, 255, 255));")
        self.delete_button.setAutoDefault(False)
        self.delete_button.setDefault(False)
        self.delete_button.setFlat(False)
        self.delete_button.setObjectName("Add")
        self.delete_button.setVisible(False)
        self.changing = QtWidgets.QPushButton(parent=MainForm)  # для изменения
        self.changing.setGeometry(QtCore.QRect(305, 560, 90, 40))
        self.changing.setStyleSheet("background-color: qlineargradient(spread:pad, x1:0.492537, y1:0.807, x2:1, y2:0, stop:0 rgba(186, 186, 186, 255), stop:1 rgba(255, 255, 255, 255));")
        self.changing.setAutoDefault(False)
        self.changing.setDefault(False)
        self.changing.setFlat(False)
        self.changing.setObjectName("changing")
        self.changing.setVisible(False)
        self.retranslateUi(MainForm)
        self.create.setShortcut("Ctrl+F1")
        self.user_inf.setShortcut("Ctrl+U")
        self.prog_inf.setShortcut("Ctrl+P")
        self.save.setShortcut("Ctrl+S")
        self.load.setShortcut("Ctrl+L")
        self.export.setShortcut("Ctrl+E")
        self.sending.setShortcut("Ctrl+F3")
        self.add_student.setShortcut("Ctrl+F2")
        self.change_student.setShortcut("Ctrl+B")
        self.delete_student.setShortcut("Ctrl+D")


        QtCore.QMetaObject.connectSlotsByName(MainForm)
        self.add_student.triggered.connect(self.creating)
        self.Cancel.clicked.connect(self.back)
        self.Add.clicked.connect(self.add)
        self.delete_student.triggered.connect(self.delete_st1)
        self.delete_button.clicked.connect(self.delete_st2)
        self.change_student.triggered.connect(self.change_st1)
        self.changing.clicked.connect(self.change_st2)
        self.export.triggered.connect(self.fileexport)
        self.save.triggered.connect(self.saveinf)
        self.load.triggered.connect(self.loadinf)
        self.sending.triggered.connect(self.file_email)
    def loadinf(self, MainForm):

        input_file = easygui.fileopenbox(filetypes=["*.txt"])
        if (input_file != None):

            self.sortspisok = []
            self.inf = []
            file= open(input_file, 'r')
            information=file.readline()
            information=information.split("|")
            for i in range(12):
                self.inf.append(information[i])
            information=file.readline()
            information=information.split("|")
            for i in range(3):
                self.count[i]=int(information[i])
            information=file.readline()
            information=information.split("|")
            for i in range(10):
                self.marks[i]=int(information[i])
            information = file.readline()
            information = information.split("|")
            for i in range(2):
                self.gpa[i] = int(information[i])
            self.tableWidget.setRowCount(self.count[0]+2)
            for i in range(self.count[0]):
                information = file.readline()
                information = information.split("|")
                self.sortspisok.append([])
                self.sortspisok[i].append(information[0])
                self.sortspisok[i].append(information[1])
                self.sortspisok[i].append(information[2])
                self.sortspisok[i].append(information[3])
                item = QtWidgets.QTableWidgetItem(str(self.sortspisok[i][0]))
                item.setFlags(QtCore.Qt.ItemFlag.ItemIsEnabled)
                self.tableWidget.setItem(i + 2, 1, QTableWidgetItem(item))
                item = QtWidgets.QTableWidgetItem(str(self.sortspisok[i][1]))
                item.setFlags(QtCore.Qt.ItemFlag.ItemIsEnabled)
                self.tableWidget.setItem(i + 2, 2, QTableWidgetItem(item))
                item = QtWidgets.QTableWidgetItem(str(self.sortspisok[i][2]))
                item.setFlags(QtCore.Qt.ItemFlag.ItemIsEnabled)
                self.tableWidget.setItem(i + 2, 3, QTableWidgetItem(item))
                item = QtWidgets.QTableWidgetItem(str(self.sortspisok[i][3]))
                item.setFlags(QtCore.Qt.ItemFlag.ItemIsEnabled)
                self.tableWidget.setItem(i + 2, 4, QTableWidgetItem(item))
            self.retranslateUi(MainForm)
            self.print_inf()
            file.close()
    def saveinf(self):
        input_file = easygui.filesavebox(filetypes=["*.txt"])
        if (input_file != None):
            input_file +=".txt"
            file= open(input_file, 'w')
            for i in range(len(self.inf)):
                file.write(str(self.inf[i]) + "|")
            file.write("\n")
            for i in range(len(self.count)):
                file.write(str(self.count[i]) + "|")
            file.write("\n")
            for i in range(len(self.marks)):
                file.write(str(self.marks[i]) + "|")
            file.write("\n")
            for i in range(len(self.gpa)):
                file.write(str(self.gpa[i]) + "|")
            file.write("\n")
            for i in range(len(self.sortspisok)):
                for j in range(len(self.sortspisok[i])):
                    file.write(str(self.sortspisok[i][j])+"|")
                file.write("\n")
            file.close()

    def file_email(self):
        input_file = "resources/шаблон.docx"
        self.file_in_word(input_file)
    def fileexport(self):
        input_file = easygui.fileopenbox(filetypes=["*.docx"])
        self.file_in_word(input_file)
    def file_in_word(self, input_file):
        gpa1 = 0
        if(input_file!=None):
            if (self.gpa[1] != 0):
                gpa1 = self.gpa[0] / self.gpa[1]
            doc = DocxTemplate(input_file)
            context = {'number': self.inf[0], 'form': self.inf[6],
                       'year': self.inf[5], 'sem': self.inf[7],
                       'facultets': self.inf[8], 'kurs': self.inf[3],
                       'group': self.inf[4], 'spec': self.inf[9],
                       'discipline': self.inf[10], 'hour': self.inf[2],
                       'teacher': self.inf[11], 'date': self.inf[1],
                       'student1': str(self.count[1]), 'student2': str(self.count[2]),
                       'ten': str(self.marks[9]), 'nine': str(self.marks[8]),
                       'eight': str(self.marks[7]), 'seven': str(self.marks[6]),
                       'six': str(self.marks[5]), 'five': str(self.marks[4]),
                       'four': str(self.marks[3]), 'three': str(self.marks[2]),
                       'two': str(self.marks[1]), 'one': str(self.marks[0]),
                       'gpa': str(round(gpa1, 2))}
            doc.render(context)
            i=input_file.find(".docx")
            input_file=input_file[:i]+"_final.docx"
            doc.save(input_file)
            doc1 = Document(input_file)
            all_tables = doc1.tables
            if(all_tables!=0):
                n=0
                for row in self.sortspisok:
                    n += 1
                    cells = doc1.tables[0].add_row().cells
                    cells[0].text = str(n)
                    for i, item in enumerate(row):
                        cells[i+1].text = str(item)

            doc1.save(input_file)


    def change_st1(self):
        if (self.count[0] == 0):
            msgBox = QMessageBox()
            msgBox.setText("В таблице нет студентов")
            msgBox.setWindowTitle("Важно")
            msgBox.setIcon(QMessageBox.Icon.Information)
            # msgBox.setStandardButtons(QMessageBox.Ok | QMessageBox.Cancel)
            # msgBox.buttonClicked.connect(msgButtonClick)
            msgBox.exec()
            return
        self.textEdit_2.setVisible(False)
        self.Cancel.setVisible(True)
        self.label_del.setVisible(True)
        self.number_stud.setVisible(True)
        self.label.setVisible(True)
        self.secondnames.setVisible(True)
        self.label_2.setVisible(True)
        self.number_book.setVisible(True)
        self.label_3.setVisible(True)
        self.ball_int.setVisible(True)
        self.label_del.setGeometry(QtCore.QRect(0, 450, 221, 16))
        self.number_stud.setGeometry(QtCore.QRect(10, 470, 111, 31))
        self.label.setGeometry(QtCore.QRect(150, 450, 221, 16))
        self.secondnames.setGeometry(QtCore.QRect(150, 470, 200, 31))
        self.label_2.setGeometry(QtCore.QRect(360, 450, 141, 16))
        self.number_book.setGeometry(QtCore.QRect(360, 470, 135, 31))
        self.label_3.setGeometry(QtCore.QRect(509, 450, 111, 16))
        self.ball_int.setGeometry(QtCore.QRect(503, 470, 111, 31))
        for i in range(self.count[0]):
            self.number_stud.addItem(str(i+1))
        self.changing.setVisible(True)

    def change_st2(self):
        self.marks_check2()
        self.sortspisok.pop(int(self.number_stud.currentText()) - 1)
        self.tableWidget.removeRow(int(self.number_stud.currentText()) + 1)
        self.count[0] -= 1#удаление
        self.count2 = self.count[0] + 2#Добаление
        self.tableWidget.setRowCount(self.count2 + 1)
        self.sortspisok.append([])
        self.sortspisok[self.count[0]].append(self.secondnames.toPlainText())
        self.sortspisok[self.count[0]].append(self.number_book.toPlainText())
        self.sortspisok[self.count[0]].append(self.ball_int.currentText())
        self.marks_check1()
        self.sortspisok.sort(key=lambda x: x[0])
        self.count[0] += 1
        for i in range(self.count[0]):
            item = QtWidgets.QTableWidgetItem(str(self.sortspisok[i][0]))
            item.setFlags(QtCore.Qt.ItemFlag.ItemIsEnabled)
            self.tableWidget.setItem(i + 2, 1, QTableWidgetItem(item))
            item = QtWidgets.QTableWidgetItem(str(self.sortspisok[i][1]))
            item.setFlags(QtCore.Qt.ItemFlag.ItemIsEnabled)
            self.tableWidget.setItem(i + 2, 2, QTableWidgetItem(item))
            item = QtWidgets.QTableWidgetItem(str(self.sortspisok[i][2]))
            item.setFlags(QtCore.Qt.ItemFlag.ItemIsEnabled)
            self.tableWidget.setItem(i + 2, 3, QTableWidgetItem(item))
            item = QtWidgets.QTableWidgetItem(str(self.sortspisok[i][3]))
            item.setFlags(QtCore.Qt.ItemFlag.ItemIsEnabled)
            self.tableWidget.setItem(i + 2, 4, QTableWidgetItem(item))
        self.secondnames.clear()
        self.number_book.clear()
        self.back()
        self.number_stud.clear()
        self.print_inf()

    def delete_st1(self):
        if(self.count[0]==0):
            msgBox = QMessageBox()
            msgBox.setText("В таблице нет студентов")
            msgBox.setWindowTitle("Важно")
            msgBox.setIcon(QMessageBox.Icon.Information)
            #msgBox.setStandardButtons(QMessageBox.Ok | QMessageBox.Cancel)
            #msgBox.buttonClicked.connect(msgButtonClick)
            msgBox.exec()
            return
        self.textEdit_2.setVisible(False)
        self.label_del.setVisible(True)
        self.Cancel.setVisible(True)
        self.number_stud.setVisible(True)
        self.delete_button.setVisible(True)
        for i in range(self.count[0]):
            self.number_stud.addItem(str(i+1))
        self.Add.setVisible(False)
        self.label.setVisible(False)
        self.label_2.setVisible(False)
        self.label_3.setVisible(False)
        self.secondnames.setVisible(False)
        self.number_book.setVisible(False)
        self.ball_int.setVisible(False)
        self.label_del.setGeometry(QtCore.QRect(260, 450, 221, 16))
        self.number_stud.setGeometry(QtCore.QRect(277, 470, 111, 31))
    def delete_st2(self):
        self.marks_check2()
        self.sortspisok.pop(int(self.number_stud.currentText())-1)
        self.tableWidget.removeRow(int(self.number_stud.currentText())+1)
        self.count[0]-=1
        self.print_inf()
        self.back()
        self.number_stud.clear()
        #self.tableWidget.selectionModel().clearCurrentIndex()

    def add(self):
        _translate = QtCore.QCoreApplication.translate
        self.count2=self.count[0]+2
        self.tableWidget.setRowCount(self.count2+1)
        self.sortspisok.append([])
        self.sortspisok[self.count[0]].append(self.secondnames.toPlainText())
        self.sortspisok[self.count[0]].append(self.number_book.toPlainText())
        self.sortspisok[self.count[0]].append(self.ball_int.currentText())
        self.marks_check1()
        self.sortspisok.sort(key = lambda x: x[0])
        self.count[0]+=1
        for i in range(self.count[0]):
            item = QtWidgets.QTableWidgetItem(str(self.sortspisok[i][0]))
            item.setFlags(QtCore.Qt.ItemFlag.ItemIsEnabled)
            self.tableWidget.setItem(i+2, 1, QTableWidgetItem(item))
            item = QtWidgets.QTableWidgetItem(str(self.sortspisok[i][1]))
            item.setFlags(QtCore.Qt.ItemFlag.ItemIsEnabled)
            self.tableWidget.setItem(i+2, 2, QTableWidgetItem(item))
            item = QtWidgets.QTableWidgetItem(str(self.sortspisok[i][2]))
            item.setFlags(QtCore.Qt.ItemFlag.ItemIsEnabled)
            self.tableWidget.setItem(i+2, 3, QTableWidgetItem(item))
            item = QtWidgets.QTableWidgetItem(str(self.sortspisok[i][3]))
            item.setFlags(QtCore.Qt.ItemFlag.ItemIsEnabled)
            self.tableWidget.setItem(i+2, 4, QTableWidgetItem(item))
        self.secondnames.clear()
        self.number_book.clear()
        self.back()
        self.print_inf()

    def marks_check1(self):
        if(self.ball_int.currentText()=="не допущен" or self.ball_int.currentText()=="не явился"):
            self.count[2]+=1
            self.sortspisok[self.count[0]].append("")

        elif(int(self.ball_int.currentText())==1):
            self.sortspisok[self.count[0]].append("один")
            self.marks[0]+=1
            self.count[1]+=1
            self.gpa[0]+=1
            self.gpa[1]+=1
        elif (int(self.ball_int.currentText()) == 2):
            self.sortspisok[self.count[0]].append("два")
            self.marks[1] += 1
            self.count[1] += 1
            self.gpa[0] += 2
            self.gpa[1] += 1
        elif (int(self.ball_int.currentText()) == 3):
            self.sortspisok[self.count[0]].append("три")
            self.marks[2] += 1
            self.count[1] += 1
            self.gpa[0] += 3
            self.gpa[1] += 1
        elif (int(self.ball_int.currentText()) == 4):
            self.sortspisok[self.count[0]].append("четыре")
            self.marks[3] += 1
            self.count[1] += 1
            self.gpa[0] += 4
            self.gpa[1] += 1
        elif (int(self.ball_int.currentText()) == 5):
            self.sortspisok[self.count[0]].append("пять")
            self.marks[4] += 1
            self.count[1] += 1
            self.gpa[0] += 5
            self.gpa[1] += 1
        elif (int(self.ball_int.currentText()) == 6):
            self.sortspisok[self.count[0]].append("шесть")
            self.marks[5] += 1
            self.count[1] += 1
            self.gpa[0] += 6
            self.gpa[1] += 1
        elif (int(self.ball_int.currentText()) == 7):
            self.sortspisok[self.count[0]].append("семь")
            self.marks[6] += 1
            self.count[1] += 1
            self.gpa[0] += 7
            self.gpa[1] += 1
        elif (int(self.ball_int.currentText()) == 8):
            self.sortspisok[self.count[0]].append("восемь")
            self.marks[7] += 1
            self.count[1] += 1
            self.gpa[0] += 8
            self.gpa[1] += 1
        elif (int(self.ball_int.currentText()) == 9):
            self.sortspisok[self.count[0]].append("девять")
            self.marks[8] += 1
            self.count[1] += 1
            self.gpa[0] += 9
            self.gpa[1] += 1
        elif (int(self.ball_int.currentText()) == 10):
            self.sortspisok[self.count[0]].append("десять")
            self.marks[9] += 1
            self.count[1] += 1
            self.gpa[0] += 10
            self.gpa[1] += 1

    def marks_check2(self):
        if(self.sortspisok[int(self.number_stud.currentText())-1][2]=="не допущен" or self.sortspisok[int(self.number_stud.currentText())-1][2]=="не явился"):
            self.count[2] -= 1
        elif(int(self.sortspisok[int(self.number_stud.currentText())-1][2])==1):
            self.count[1] -= 1
            self.gpa[1] -= 1
            self.gpa[0] -= int(self.sortspisok[int(self.number_stud.currentText()) - 1][2])
            self.marks[0]-=1
        elif (int(self.sortspisok[int(self.number_stud.currentText())-1][2]) == 2):
            self.count[1] -= 1
            self.gpa[1] -= 1
            self.gpa[0] -= int(self.sortspisok[int(self.number_stud.currentText()) - 1][2])
            self.marks[1] -= 1
        elif (int(self.sortspisok[int(self.number_stud.currentText())-1][2]) == 3):
            self.count[1] -= 1
            self.gpa[1] -= 1
            self.gpa[0] -= int(self.sortspisok[int(self.number_stud.currentText()) - 1][2])
            self.marks[2] -= 1
        elif (int(self.sortspisok[int(self.number_stud.currentText())-1][2]) == 4):
            self.count[1] -= 1
            self.gpa[1] -= 1
            self.gpa[0] -= int(self.sortspisok[int(self.number_stud.currentText()) - 1][2])
            self.marks[3] -= 1
        elif (int(self.sortspisok[int(self.number_stud.currentText())-1][2]) == 5):
            self.count[1] -= 1
            self.gpa[1] -= 1
            self.gpa[0] -= int(self.sortspisok[int(self.number_stud.currentText()) - 1][2])
            self.marks[4] -= 1
        elif (int(self.sortspisok[int(self.number_stud.currentText())-1][2]) == 6):
            self.count[1] -= 1
            self.gpa[1] -= 1
            self.gpa[0] -= int(self.sortspisok[int(self.number_stud.currentText()) - 1][2])
            self.marks[5] -= 1
        elif (int(self.sortspisok[int(self.number_stud.currentText())-1][2]) == 7):
            self.count[1] -= 1
            self.gpa[1] -= 1
            self.gpa[0] -= int(self.sortspisok[int(self.number_stud.currentText()) - 1][2])
            self.marks[6] -= 1
        elif (int(self.sortspisok[int(self.number_stud.currentText())-1][2]) == 8):
            self.count[1] -= 1
            self.gpa[1] -= 1
            self.gpa[0] -= int(self.sortspisok[int(self.number_stud.currentText()) - 1][2])
            self.marks[7] -= 1
        elif (int(self.sortspisok[int(self.number_stud.currentText())-1][2]) == 9):
            self.count[1] -= 1
            self.gpa[1] -= 1
            self.gpa[0] -= int(self.sortspisok[int(self.number_stud.currentText()) - 1][2])
            self.marks[8] -= 1
        elif (int(self.sortspisok[int(self.number_stud.currentText())-1][2]) == 10):
            self.count[1] -= 1
            self.gpa[1] -= 1
            self.gpa[0] -= int(self.sortspisok[int(self.number_stud.currentText()) - 1][2])
            self.marks[9] -= 1

    def creating(self, inf):
        self.textEdit_2.setVisible(False)
        self.Add.setVisible(True)
        self.Cancel.setVisible(True)
        self.label.setVisible(True)
        self.label_2.setVisible(True)
        self.label_3.setVisible(True)
        self.secondnames.setVisible(True)
        self.number_book.setVisible(True)
        self.ball_int.setVisible(True)
        self.label_del.setVisible(False)
        self.number_stud.setVisible(False)
        self.delete_button.setVisible(False)
        self.label.setGeometry(QtCore.QRect(0, 450, 221, 16))
        self.secondnames.setGeometry(QtCore.QRect(0, 470, 211, 31))
        self.label_2.setGeometry(QtCore.QRect(225, 450, 141, 16))
        self.number_book.setGeometry(QtCore.QRect(225, 470, 141, 31))
        self.label_3.setGeometry(QtCore.QRect(370, 450, 111, 16))
        self.ball_int.setGeometry(QtCore.QRect(370, 470, 111, 31))

    def back(self):
        self.textEdit_2.setVisible(True)
        self.Add.setVisible(False)
        self.Cancel.setVisible(False)
        self.label.setVisible(False)
        self.label_2.setVisible(False)
        self.label_3.setVisible(False)
        self.secondnames.setVisible(False)
        self.number_book.setVisible(False)
        self.ball_int.setVisible(False)
        self.label_del.setVisible(False)
        self.number_stud.setVisible(False)
        self.delete_button.setVisible(False)
        self.changing.setVisible(False)

    def retranslateUi(self, MainForm):
        _translate = QtCore.QCoreApplication.translate

        __sortingEnabled = self.tableWidget.isSortingEnabled()
        self.tableWidget.setSortingEnabled(False)
        item = self.tableWidget.item(0, 0)
        item.setText(_translate("MainForm", " № п/п"))
        item = self.tableWidget.item(0, 1)
        item.setText(_translate("MainForm", "Фамилия, инициалы обучащаегося"))
        item = self.tableWidget.item(0, 2)
        item.setText(_translate("MainForm", "№ зачётной книжки"))
        item = self.tableWidget.item(0, 3)
        item.setText(_translate("MainForm", "Отметка в баллах"))
        item = self.tableWidget.item(0, 4)
        item.setText(_translate("MainForm", "Отметка прописью"))
        self.tableWidget.setSortingEnabled(__sortingEnabled)
        self.textEdit.setHtml(_translate("MainForm", "<!DOCTYPE HTML PUBLIC \"-//W3C//DTD HTML 4.0//EN\" \"http://www.w3.org/TR/REC-html40/strict.dtd\">\n"
"<html><head><meta name=\"qrichtext\" content=\"1\" /><style type=\"text/css\">\n"
"p, li { white-space: pre-wrap; }\n"
"</style></head><body style=\" font-family:\'MS Shell Dlg 2\'; font-size:7.8pt; font-weight:400; font-style:normal;\">\n"
"<p align=\"center\" style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><span style=\" font-size:10pt;\">Экзамиционная ведомость №"+self.inf[0]+"</span></p>\n"
"<p style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><span style=\" font-size:10pt;\">Форма получения "+self.inf[6]+"</span></p>\n"
"<p style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><span style=\" font-size:10pt;\">Учебный год "+self.inf[5]+", семестр "+self.inf[7]+"</span></p>\n"
"<p style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><span style=\" font-size:10pt;\">Факультет "+self.inf[8]+"</span></p>\n"
"<p style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><span style=\" font-size:10pt;\">Курс "+self.inf[3]+", группа "+self.inf[4]+"</span></p>\n"
"<p style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><span style=\" font-size:10pt;\">Специальность "+self.inf[9]+"</span></p>\n"
"<p style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><span style=\" font-size:10pt;\">Дисциплина "+self.inf[10]+"</span></p>\n"
"<p style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><span style=\" font-size:10pt;\">Всего часов по дисциплине в семестре "+self.inf[2]+"</span></p>\n"
"<p style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><span style=\" font-size:10pt;\">Фамилия, инициалы преподователя(ей) "+self.inf[11]+"</span></p>\n"
"<p style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><span style=\" font-size:10pt;\">Дата проведения аттестации "+self.inf[1]+"</span></p>\n"
"<p style=\"-qt-paragraph-type:empty; margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px; font-size:12pt;\"><br /></p></body></html>"))
        self.textEdit_2.setHtml(_translate("MainForm", "<!DOCTYPE HTML PUBLIC \"-//W3C//DTD HTML 4.0//EN\" \"http://www.w3.org/TR/REC-html40/strict.dtd\">\n"
"<html><head><meta name=\"qrichtext\" content=\"1\" /><style type=\"text/css\">\n"
"p, li { white-space: pre-wrap; }\n"
"</style></head><body style=\" font-family:\'MS Shell Dlg 2\'; font-size:7.8pt; font-weight:400; font-style:normal;\">\n"
"<p style=\" margin-top:12px; margin-bottom:12px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><span style=\" font-size:10pt;\">Количество присутсвующих на аттестации: "+str(self.count[0])+"</span></p>\n"
"<p style=\" margin-top:12px; margin-bottom:12px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><span style=\" font-size:10pt;\">Количество получивших отметки:</span></p>\n"
"<p style=\" margin-top:12px; margin-bottom:12px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><span style=\" font-size:10pt;\">10(десять): "+str(self.marks[9])+"            9(девять): "+str(self.marks[8])+"            8(восемь): "+str(self.marks[7])+"              7(семь): "+str(self.marks[6])+"                  6(шесть): "+str(self.marks[5])+"         5(пять): "+str(self.marks[4])+"</span>                        <span style=\" font-size:10pt;\">4(четыре): "+str(self.marks[3])+"             3(три): "+str(self.marks[2])+"                 2(два): "+str(self.marks[1])+"                   1(один): "+str(self.marks[0])+"                  Средний балл: 0</span>                  <span style=\" font-size:10pt;\">        </span></p>\n"
"<p style=\" margin-top:12px; margin-bottom:12px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><span style=\" font-size:10pt;\">Количество обучающихся не явивщихся на аттестацию(в том числе и не допущенные): "+str(self.count[2])+"</span></p></body></html>"))
        self.menu.setTitle(_translate("MainForm", "Файл"))
        self.menu_3.setTitle(_translate("MainForm", "Работа"))
        self.menu_2.setTitle(_translate("MainForm", "Информация"))
        self.create.setText(_translate("MainForm", "Создать"))
        self.load.setText(_translate("MainForm", "Загрузить"))
        self.save.setText(_translate("MainForm", "Сохранить"))
        self.export.setText(_translate("MainForm", "Экспорт в Word"))
        self.sending.setText(_translate("MainForm", "Отправка по Email"))
        self.user_inf.setText(_translate("MainForm", "Об авторе"))
        self.prog_inf.setText(_translate("MainForm", "О программе"))
        self.add_student.setText(_translate("MainForm", "Добавить"))
        self.change_student.setText(_translate("MainForm", "Изменить"))
        self.delete_student.setText(_translate("MainForm", "Удалить"))
        self.Add.setText(_translate("MainForm", "Добавить"))#Добавление
        self.Cancel.setText(_translate("MainForm", "Назад"))
        self.label.setText(_translate("MainForm", "Фамилия и инициалы обучаещегося"))
        self.label_2.setText(_translate("MainForm", "Номер зачетной книжки"))
        self.label_3.setText(_translate("MainForm", "Отметка в баллах"))
        self.ball_int.setItemText(0, _translate("MainForm", "1"))
        self.ball_int.setItemText(1, _translate("MainForm", "2"))
        self.ball_int.setItemText(2, _translate("MainForm", "3"))
        self.ball_int.setItemText(3, _translate("MainForm", "4"))
        self.ball_int.setItemText(4, _translate("MainForm", "5"))
        self.ball_int.setItemText(5, _translate("MainForm", "6"))
        self.ball_int.setItemText(6, _translate("MainForm", "7"))
        self.ball_int.setItemText(7, _translate("MainForm", "8"))
        self.ball_int.setItemText(8, _translate("MainForm", "9"))
        self.ball_int.setItemText(9, _translate("MainForm", "10"))
        self.ball_int.setItemText(10, _translate("MainForm", "не допущен"))
        self.ball_int.setItemText(11, _translate("MainForm", "не явился"))
        self.label.setText(_translate("Add_Student", "Фамилия и инициалы обучаещегося"))#удаление
        self.label_del.setText(_translate("Add_Student", "Выберите номер студента"))
        self.delete_button.setText(_translate("MainForm", "Удалить"))
        self.changing.setText(_translate("MainForm", "Изменить"))#изменение

    def print_inf(self):
        rows = self.tableWidget.rowCount()
        columns = self.tableWidget.columnCount()

        for i in range(2,self.count[0]+2):
            item = QTableWidgetItem(str(i-1))
            item.setTextAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
            item.setFlags(QtCore.Qt.ItemFlag.ItemIsEnabled)
            self.tableWidget.setItem(i, 0, item)
        gpa1=0
        if(self.gpa[1]!=0):
            gpa1=self.gpa[0]/self.gpa[1]
        _translate = QtCore.QCoreApplication.translate
        self.textEdit_2.setHtml(_translate("MainForm",
            "<!DOCTYPE HTML PUBLIC \"-//W3C//DTD HTML 4.0//EN\" \"http://www.w3.org/TR/REC-html40/strict.dtd\">\n"
            "<html><head><meta name=\"qrichtext\" content=\"1\" /><style type=\"text/css\">\n"
            "p, li { white-space: pre-wrap; }\n"
            "</style></head><body style=\" font-family:\'MS Shell Dlg 2\'; font-size:7.8pt; font-weight:400; font-style:normal;\">\n"
            "<p style=\" margin-top:12px; margin-bottom:12px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><span style=\" font-size:10pt;\">Количество присутсвующих на аттестации " + str(self.count[1]) + "</span></p>\n"
            "<p style=\" margin-top:12px; margin-bottom:12px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><span style=\" font-size:10pt;\">Количество получивших отметки</span></p>\n"
            "<p style=\" margin-top:12px; margin-bottom:12px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><span style=\" font-size:10pt;\">10(десять): "+str(self.marks[9])+"            9(девять): "+str(self.marks[8])+"            8(восемь): "+str(self.marks[7])+"              7(семь): "+str(self.marks[6])+"                  6(шесть): "+str(self.marks[5])+"         5(пять): "+str(self.marks[4])+"</span>                        <span style=\" font-size:10pt;\">4(четыре): "+str(self.marks[3])+"             3(три): "+str(self.marks[2])+"                 2(два): "+str(self.marks[1])+"                   1(один): "+str(self.marks[0])+"                  Средний балл:" +str(round(gpa1, 2))+"</span>                  <span style=\" font-size:10pt;\">        </span></p>\n"
            "<p style=\" margin-top:12px; margin-bottom:12px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><span style=\" font-size:10pt;\">Количество обучающихся не явивщихся на аттестацию(в том числе и не допущенные): "+str(self.count[2])+"</span></p></body></html>"))

